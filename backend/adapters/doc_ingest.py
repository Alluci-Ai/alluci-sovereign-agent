
import os
import re
import hashlib
from ..logging_config import get_logger
from typing import Dict, Any, List, Optional
try:
    from pypdf import PdfReader
except ImportError:
    class PdfReader:
        def __init__(self, file_path):
            self.pages = []  # No pages; stub

try:
    from docx import Document
except ImportError:
    class Document:
        def __init__(self, file_path=None):
            self.paragraphs = []  # Empty list of paragraphs

from .base import Adapter
from ..memory.manager import MemoryManager
from ..engine.vpi import VisualPolytopeIngestor

FIGURE_CAPTION_REGEX = re.compile(
    r'(?i)(?:Figure|Fig\.|Chart|Diagram|Graph|Illustration|Schematic)\s+(\d+(?:\.\d+)*)[:\.\-–—]\s*([^\n\r]+)'
)

class DocumentIngestAdapter(Adapter):
    """
    Document Ingestion Adapter.
    Extracts text, structural pages, and substantive technical figures (diagrams, flowcharts, charts),
    filters decorative visual noise, and indexes entities across 4-Tier H-LSM memory.
    """
    name = "doc_ingest"
    description = "Ingest and index PDF, DOCX, or TXT documents along with substantive technical figures into memory."

    def __init__(self, memory_manager: MemoryManager):
        self.memory_manager = memory_manager
        self.logger = get_logger("DocumentIngestAdapter")
        self.vpi = VisualPolytopeIngestor()

    async def execute(self, args: Dict[str, Any]) -> Dict[str, Any]:
        """
        Reads a document, extracts text and technical figures, chunks text,
        and stores structured entities in long-term memory.

        args:
            file_path (str): Absolute or relative path to the file to ingest.
            session_key (str): Optional session key.
        """
        file_path = args.get("file_path") or args.get("path") if isinstance(args, dict) else str(args)
        session_key = args.get("session_key", "") if isinstance(args, dict) else ""

        if not file_path:
            return {"status": "error", "message": "No 'file_path' provided in args."}

        try:
            if not os.path.exists(file_path):
                return {"status": "error", "message": f"File not found: {file_path}"}

            ext = os.path.splitext(file_path)[1].lower()
            filename = os.path.basename(file_path)
            doc_slug = re.sub(r'[^a-zA-Z0-9_]', '_', os.path.splitext(filename)[0]).lower()
            text = ""
            raw_candidate_figures: List[Dict[str, Any]] = []

            # Create figure output directory under workspace/artifacts/extracted_figures/<doc_slug>/
            fig_output_dir = os.path.join("workspace", "artifacts", "extracted_figures", doc_slug)
            os.makedirs(fig_output_dir, exist_ok=True)

            if ext == ".pdf":
                reader = PdfReader(file_path)
                for p_idx, page in enumerate(reader.pages):
                    page_num = p_idx + 1
                    page_text = page.extract_text() or ""
                    text += f"\n--- [DOCUMENT: {filename} | PAGE {page_num}/{len(reader.pages)}] ---\n" + page_text + "\n"

                    # Caption extraction from page text
                    captions = FIGURE_CAPTION_REGEX.findall(page_text)

                    # Image extraction from PDF page
                    if hasattr(page, "images"):
                        for img_idx, img in enumerate(page.images):
                            try:
                                img_data = img.data
                                img_name = getattr(img, "name", f"img_p{page_num}_{img_idx}.png")
                                img_sha = hashlib.sha256(img_data).hexdigest()
                                fig_filename = f"fig_p{page_num}_{img_idx}_{img_sha[:8]}.png"
                                fig_path = os.path.join(fig_output_dir, fig_filename)

                                with open(fig_path, "wb") as f_img:
                                    f_img.write(img_data)

                                # Associate nearest caption if available
                                matched_caption = ""
                                if img_idx < len(captions):
                                    c_num, c_text = captions[img_idx]
                                    matched_caption = f"Figure {c_num}: {c_text.strip()}"
                                elif captions:
                                    c_num, c_text = captions[0]
                                    matched_caption = f"Figure {c_num}: {c_text.strip()}"

                                raw_candidate_figures.append({
                                    "id": f"fig_{doc_slug[:12]}_p{page_num}_{img_idx}",
                                    "document_id": doc_slug,
                                    "page_number": page_num,
                                    "caption": matched_caption,
                                    "file_path": fig_path,
                                    "is_vector": False,
                                    "sha256": img_sha,
                                    "width": getattr(img, "width", 800),
                                    "height": getattr(img, "height", 600),
                                    "extracted_text": matched_caption,
                                })
                            except Exception as img_err:
                                self.logger.debug(f"[DOC_INGEST] Image extraction notice on p{page_num}: {img_err}")

            elif ext == ".docx":
                doc = Document(file_path)
                for para in doc.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"
            elif ext in [".txt", ".md"]:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            else:
                return {"status": "error", "message": f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt, .md"}

            if not text.strip():
                return {"status": "error", "message": f"No extractable text found in {file_path}"}

            # Filter candidate figures using Dual-Pass VPI evaluation
            substantive_figures: List[Dict[str, Any]] = []
            if raw_candidate_figures:
                substantive_figures = self.vpi.filter_and_caption_figures(
                    raw_candidate_figures,
                    document_id=doc_slug
                )

            # Route through H-LSM Manager if available
            from .. import services
            if hasattr(services, "hlsm_manager") and services.hlsm_manager is not None:
                await services.hlsm_manager.ingest_document_payload(
                    filename=filename,
                    content=text,
                    session_key=session_key,
                    metadata={
                        "file_path": file_path,
                        "mime_type": f"application/{ext[1:]}",
                        "figures": substantive_figures
                    }
                )

            chunks = self._chunk_text(text, ext=ext)

            for i, chunk in enumerate(chunks):
                await self.memory_manager.store(
                    content=chunk,
                    metadata={
                        "source": file_path,
                        "filename": filename,
                        "file_type": ext[1:],
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "figures_count": len(substantive_figures),
                    },
                )
                self._extract_entities_for_graph(chunk)

            self.logger.info(f"[ DOC_INGEST ] Ingested {len(chunks)} chunks and {len(substantive_figures)} figures from '{filename}'")
            return {
                "status": "success",
                "message": f"Ingested {len(chunks)} chunks and {len(substantive_figures)} technical figures from {filename}",
                "chunks_count": len(chunks),
                "figures_count": len(substantive_figures),
                "figures": substantive_figures,
                "file_path": file_path,
            }
        except Exception as e:
            self.logger.error(f"Document ingestion failed for '{file_path}': {e}")
            return {"status": "error", "message": str(e)}

    def _chunk_text(self, text: str, chunk_size: int = 1000, ext: str = ".txt") -> List[str]:
        """
        Splits text into manageable chunks.
        For markdown files, it attempts to split by headers semantically.
        """
        if ext == ".md":
            chunks = re.split(r'(?m)^#{1,6}\s+.*$', text)
            headers = re.findall(r'(?m)^#{1,6}\s+.*$', text)
            
            result = []
            if chunks and chunks[0].strip():
                result.append(chunks[0].strip())
            
            for i, header in enumerate(headers):
                if i + 1 < len(chunks):
                    chunk_content = header + "\n" + chunks[i + 1].strip()
                    if chunk_content.strip():
                        result.append(chunk_content.strip())
            
            final_chunks = []
            for c in result:
                if len(c) > chunk_size * 2:
                    final_chunks.extend([c[i:i + chunk_size] for i in range(0, len(c), chunk_size)])
                else:
                    final_chunks.append(c)
            if final_chunks:
                return final_chunks
                
        return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

    def _extract_entities_for_graph(self, chunk: str) -> None:
        """
        Extracts Graph entities (Nodes/Edges) for L3 Kuzu ingestion.
        Implemented via cognitive LLM pass in the broader ingestion pipeline.
        """
        pass

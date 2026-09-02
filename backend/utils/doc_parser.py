"""
Universal Document & Media Parsing Engine
==========================================

Extracts clean, human-readable text and structured metadata from multi-format files:
- Documents: PDF, TXT, MD, RTF
- Word Documents: DOCX
- Code & Configs: PY, TS, JS, JSON, YAML, XML, HTML, SH, CSS
- Data Tables: CSV, TSV
- Vision Media: PNG, JPG, JPEG, WEBP
- Audio Media: MP3, WAV, M4A
"""

import os
import io
import re
import csv
import json
import base64
import zipfile
import hashlib
import xml.etree.ElementTree as ET
from html.parser import HTMLParser
from typing import Dict, Any, List, Optional, Tuple

from ..logging_config import get_logger
from ..engine.vpi import VisualPolytopeIngestor

logger = get_logger("DocParser")

FIGURE_CAPTION_REGEX = re.compile(
    r'(?i)(?:Figure|Fig\.|Chart|Diagram|Graph|Illustration|Schematic)\s+(\d+(?:\.\d+)*)[:\.\-–—]\s*([^\n\r]+)'
)


class MLStripper(HTMLParser):
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text = []
        self.ignore = False

    def handle_starttag(self, tag, attrs):
        if tag in ["style", "script", "head", "title", "meta"]:
            self.ignore = True

    def handle_endtag(self, tag):
        if tag in ["style", "script", "head", "title", "meta"]:
            self.ignore = False

    def handle_data(self, data: str):
        if not self.ignore:
            self.text.append(data)

    def get_data(self):
        return ''.join(self.text)


def strip_html_tags(html_content: str) -> str:
    """Strips HTML/XML tags cleanly while preserving text paragraphs."""
    if not html_content or ('<' not in html_content and '>' not in html_content):
        return html_content
    try:
        stripper = MLStripper()
        stripper.feed(html_content)
        cleaned = re.sub(r'\n\s*\n', '\n\n', stripper.get_data())
        return cleaned.strip()
    except Exception:
        cleaned = re.sub(r'<[^>]+>', '', html_content)
        return re.sub(r'\n\s*\n', '\n\n', cleaned).strip()


def strip_rtf_formatting(rtf_text: str) -> str:
    """Strips RTF control tags and returns plain text content."""
    pattern = r'\\[a-z0-9]+(?:\s|-?\d+)?|\\\'[0-9a-f]{2}|\{|\}'
    cleaned = re.sub(pattern, '', rtf_text, flags=re.IGNORECASE)
    cleaned = re.sub(r'\n\s*\n', '\n\n', cleaned)
    return cleaned.strip()


def extract_pages_from_pdf_bytes(pdf_bytes: bytes, filename: str = "") -> List[Dict[str, Any]]:
    """Extracts structured page objects from PDF raw bytes using pypdf with regex fallback."""
    pages_list = []
    fname = os.path.basename(filename) if filename else "document.pdf"
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        total_pages = len(reader.pages)
        for idx, page in enumerate(reader.pages):
            p_num = idx + 1
            t = page.extract_text() or ""
            clean_text = t.strip()
            pages_list.append({
                "page_number": p_num,
                "total_pages": total_pages,
                "text": clean_text,
                "char_count": len(clean_text),
                "filename": fname,
                "header": f"--- [DOCUMENT: {fname} | PAGE {p_num}/{total_pages}] ---"
            })
        if pages_list:
            return pages_list
    except Exception as pdf_err:
        logger.warning(f"[DocParser] pypdf page extraction notice: {pdf_err}")

    # Regex stream fallback (treated as page 1)
    try:
        matches = re.findall(rb'\((.*?)\)', pdf_bytes)
        clean_strings = [
            m.decode('utf-8', errors='ignore').strip() 
            for m in matches 
            if len(m) > 2 and any(c.isalnum() for c in m.decode('utf-8', errors='ignore'))
        ]
        if clean_strings:
            combined = " ".join(clean_strings)
            return [{
                "page_number": 1,
                "total_pages": 1,
                "text": combined,
                "char_count": len(combined),
                "filename": fname,
                "header": f"--- [DOCUMENT: {fname} | PAGE 1/1] ---"
            }]
    except Exception:
        pass

    return []


def extract_text_from_pdf_bytes(pdf_bytes: bytes, filename: str = "") -> str:
    """Extracts text pages from PDF raw bytes with explicit page boundary headers."""
    pages = extract_pages_from_pdf_bytes(pdf_bytes, filename=filename)
    if not pages:
        return ""
    
    formatted_pages = []
    for p in pages:
        p_text = p.get("text", "").strip()
        p_hdr = p.get("header", "")
        if p_text:
            formatted_pages.append(f"{p_hdr}\n{p_text}")
        else:
            formatted_pages.append(f"{p_hdr}\n[Empty page or image-only content]")
            
    return "\n\n".join(formatted_pages)


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Extracts paragraphs and tables from Word (.docx) raw bytes using python-docx with stdlib zipfile fallback."""
    text_chunks = []
    # Method 1: python-docx
    try:
        import docx
        doc = docx.Document(io.BytesIO(docx_bytes))
        for p in doc.paragraphs:
            if p.text.strip():
                text_chunks.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_chunks.append(" | ".join(row_text))
        if text_chunks:
            return "\n\n".join(text_chunks)
    except Exception as docx_err:
        logger.debug(f"[DocParser] python-docx extraction notice: {docx_err}")

    # Method 2: stdlib zipfile + xml.etree.ElementTree fallback
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            if "word/document.xml" in zf.namelist():
                xml_content = zf.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                texts = [node.text for node in tree.iter() if node.tag.endswith('t') and node.text]
                if texts:
                    return "\n".join(texts)
    except Exception as zip_err:
        logger.debug(f"[DocParser] zipfile docx extraction notice: {zip_err}")

    return ""


def extract_text_from_csv_bytes(csv_bytes: bytes, delimiter: str = ",") -> str:
    """Parses CSV/TSV data into Markdown tables."""
    try:
        decoded = csv_bytes.decode("utf-8", errors="replace")
        lines = [line for line in decoded.splitlines() if line.strip()]
        if not lines:
            return ""

        reader = csv.reader(lines, delimiter=delimiter)
        rows = list(reader)
        if not rows:
            return ""

        md_lines = []
        header = rows[0]
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")

        for r in rows[1:100]:  # Limit to first 100 rows for context sanity
            md_lines.append("| " + " | ".join(r) + " |")

        if len(rows) > 101:
            md_lines.append(f"\n*... ({len(rows) - 101:,} additional rows truncated for context safety) ...*")

        return "\n".join(md_lines)
    except Exception as e:
        logger.warning(f"[DocParser] CSV parsing error: {e}")
        return csv_bytes.decode("utf-8", errors="replace")


def extract_document_and_figures_from_payload(
    file_name: str, 
    file_data_base64: str, 
    file_mime: str = ""
) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Unified multimodal document and figure extraction dispatcher.
    Extracts text and substantive technical figures (diagrams, flowcharts, data graphs),
    filters decorative visual noise using VPI dual-pass evaluation, and persists figures to
    workspace/artifacts/extracted_figures/<doc_slug>/.
    
    Returns:
        (extracted_text, substantive_figures)
    """
    if not file_data_base64:
        return "", []

    fn_lower = file_name.lower()
    mime_lower = file_mime.lower()

    try:
        raw_bytes = base64.b64decode(file_data_base64)
    except Exception as e:
        logger.warning(f"[DocParser] Failed to decode base64 for {file_name}: {e}")
        return file_data_base64, []

    doc_slug = re.sub(r'[^a-zA-Z0-9_]', '_', os.path.splitext(os.path.basename(file_name))[0]).lower()

    # 1. PDF Documents (Multimodal extraction: text + raster figures)
    is_pdf = "pdf" in mime_lower or fn_lower.endswith(".pdf") or raw_bytes.startswith(b"%PDF")
    if is_pdf:
        extracted_text = ""
        raw_candidate_figures: List[Dict[str, Any]] = []
        fig_output_dir = os.path.join("workspace", "artifacts", "extracted_figures", doc_slug)
        os.makedirs(fig_output_dir, exist_ok=True)

        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
            total_pages = len(reader.pages)
            for p_idx, page in enumerate(reader.pages):
                page_num = p_idx + 1
                page_text = page.extract_text() or ""
                extracted_text += f"\n--- [DOCUMENT: {file_name} | PAGE {page_num}/{total_pages}] ---\n{page_text}\n"

                captions = FIGURE_CAPTION_REGEX.findall(page_text)

                if hasattr(page, "images"):
                    for img_idx, img in enumerate(page.images):
                        try:
                            img_data = img.data
                            img_sha = hashlib.sha256(img_data).hexdigest()
                            fig_filename = f"fig_p{page_num}_{img_idx}_{img_sha[:8]}.png"
                            fig_path = os.path.join(fig_output_dir, fig_filename)

                            with open(fig_path, "wb") as f_img:
                                f_img.write(img_data)

                            matched_caption = ""
                            if img_idx < len(captions):
                                c_num, c_text = captions[img_idx]
                                matched_caption = f"Figure {c_num}: {c_text.strip()}"
                            elif captions:
                                c_num, c_text = captions[0]
                                matched_caption = f"Figure {c_num}: {c_text.strip()}"

                            raw_candidate_figures.append({
                                "id": f"fig_{doc_slug[:12]}_p{page_num}_{img_idx}",
                                "document_id": doc_slug,
                                "page_number": page_num,
                                "caption": matched_caption,
                                "file_path": fig_path,
                                "is_vector": False,
                                "sha256": img_sha,
                                "width": getattr(img, "width", 800),
                                "height": getattr(img, "height", 600),
                                "extracted_text": matched_caption,
                            })
                        except Exception as img_err:
                            logger.debug(f"[DocParser] PDF image extraction notice on p{page_num}: {img_err}")
        except Exception as pdf_err:
            logger.warning(f"[DocParser] pypdf multimodal extraction error: {pdf_err}")
            extracted_text = extract_text_from_pdf_bytes(raw_bytes, filename=file_name)

        # Filter candidate figures using Dual-Pass VPI evaluation
        substantive_figures: List[Dict[str, Any]] = []
        if raw_candidate_figures:
            vpi = VisualPolytopeIngestor()
            substantive_figures = vpi.filter_and_caption_figures(
                raw_candidate_figures,
                document_id=doc_slug
            )

        if not extracted_text.strip():
            extracted_text = f"[PDF DOCUMENT: {file_name} — No readable text stream extracted]"

        return extracted_text.strip(), substantive_figures

    # 2. Word Documents (.docx)
    is_docx = "wordprocessingml" in mime_lower or fn_lower.endswith(".docx") or (raw_bytes.startswith(b"PK\x03\x04") and fn_lower.endswith((".docx", ".doc")))
    if is_docx:
        docx_text = extract_text_from_docx_bytes(raw_bytes)
        if docx_text.strip():
            return docx_text.strip(), []
        return f"[WORD DOCUMENT: {file_name} — No extractable paragraph text found]", []

    # 3. CSV / TSV Data Tables
    if fn_lower.endswith(".csv") or "csv" in mime_lower:
        return extract_text_from_csv_bytes(raw_bytes, delimiter=","), []
    if fn_lower.endswith(".tsv") or "tab-separated" in mime_lower:
        return extract_text_from_csv_bytes(raw_bytes, delimiter="\t"), []

    # 4. RTF Documents
    if fn_lower.endswith(".rtf") or "rtf" in mime_lower:
        try:
            raw_text = raw_bytes.decode("utf-8", errors="replace")
            return strip_rtf_formatting(raw_text), []
        except Exception:
            pass

    # 5. Image & Media Formats
    is_image = any(ext in fn_lower for ext in [".png", ".jpg", ".jpeg", ".webp"]) or "image/" in mime_lower
    if is_image:
        fig_output_dir = os.path.join("workspace", "artifacts", "extracted_figures", doc_slug)
        os.makedirs(fig_output_dir, exist_ok=True)
        img_sha = hashlib.sha256(raw_bytes).hexdigest()
        fig_filename = f"img_direct_{img_sha[:8]}_{os.path.basename(file_name)}"
        fig_path = os.path.join(fig_output_dir, fig_filename)
        with open(fig_path, "wb") as f_img:
            f_img.write(raw_bytes)
        
        fig_obj = {
            "id": f"fig_{doc_slug[:12]}_{img_sha[:8]}",
            "document_id": doc_slug,
            "page_number": 1,
            "caption": file_name,
            "file_path": fig_path,
            "is_vector": False,
            "sha256": img_sha,
            "figure_type": "TECHNICAL_FIGURE",
            "visual_summary": f"Uploaded visual image asset: {file_name}",
            "is_substantive": True
        }
        return f"[IMAGE ATTACHMENT: {file_name} ({len(raw_bytes):,} bytes) — Image payload saved for Vision Grounding]", [fig_obj]

    # 6. Audio Media Formats
    is_audio = any(ext in fn_lower for ext in [".mp3", ".wav", ".m4a"]) or "audio/" in mime_lower
    if is_audio:
        return f"[AUDIO ATTACHMENT: {file_name} ({len(raw_bytes):,} bytes) — Audio track ready for Whisper Transcription]", []

    # 7. Code, Configs, HTML, and Plain Text
    try:
        decoded_text = raw_bytes.decode("utf-8")
        if decoded_text.startswith(("PK\x03\x04", "\x7fELF", "MZ")):
            return f"[BINARY FILE: {file_name} — Binary format cannot be rendered as text]", []
        if fn_lower.endswith((".html", ".htm", ".xml")):
            return strip_html_tags(decoded_text), []
        return decoded_text.strip(), []
    except UnicodeDecodeError:
        pass

    try:
        return raw_bytes.decode("latin-1", errors="replace").strip(), []
    except Exception:
        return f"[UNSUPPORTED FILE FORMAT: {file_name}]", []


def extract_text_from_file_payload(file_name: str, file_data_base64: str, file_mime: str = "") -> str:
    """
    Unified multi-format file extraction dispatcher (Text-only string output).
    Maintained for backward compatibility.
    """
    text, _ = extract_document_and_figures_from_payload(file_name, file_data_base64, file_mime)
    return text
